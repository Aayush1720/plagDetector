from django.http import JsonResponse, HttpResponse
from django.shortcuts import render, redirect
from .forms import UploadFileForm
from django.core.files.storage import FileSystemStorage
from auth_app.models import User
from .algorithms import plagiarism_detector
import os
from .document_parser import document_parser

from docx import Document
from docx.shared import Inches, Cm
import statistics
from pylab import rcParams
import seaborn as sns
import matplotlib
import matplotlib.pyplot as plt


import smtplib, ssl
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
import secrets

#########################################################


def home(request):
    context = {'navigation': []}
    is_authenticated = request.user.is_authenticated
    notAuth = not is_authenticated
    username = ""
    if is_authenticated:
        username = request.user.username

    Matrix = [['x', 'a', 'b', 'c'], ['a', 1, 3, 5], ['b', 4, 7, 34], ['c', 32, 8, 43]]
    context = {'is_authenticated': is_authenticated, 'username': username, "notAuth": notAuth, 'result': Matrix, }
    return render(request, 'plagDetector/home.html', context)


# test function to compare two texts -Tanmay/Aadarsh/Aayush
def get_similarity(doc_list):
    # str1 = ''.join(text1)
    # str2 = ''.join(text2)
    # doc_list = [str1, str2]
    pd = plagiarism_detector(doc_list)
    similarity = pd.get_results()
    return similarity


# test function two check two input texts
def search(request):
    # search goes here
    if request.method == 'POST':
        final_list = "cosine result"
        text1 = request.POST.get('text1', None)
        text2 = request.POST.get('text2', None)
        result = get_similarity([text1, text2])
        is_authenticated = request.user.is_authenticated
        username = ""
        if is_authenticated:
            username = request.user.username
        context = {'result': result, 'is_authenticated': is_authenticated, 'username': username}
        print(context)
        return render(request, 'plagDetector/home.html', context)
    else:
        return render(request, 'plagDetector/home.html')


def generate_email(request, user_ass):
    email = request.user.email
    print(email)
    print(user_ass)
    print(request.path)
    print(request.get_full_path)
    # return redirect("upload")
    print("sending email.....")
    smtp_server = 'smtp.gmail.com'
    smtp_port = 587
    # Replace with your own gmail account
    EMAIL = "tanmaymodi9@gmail.com"
    PASSWORD = "Password@123"
    gmail = EMAIL
    password = PASSWORD

    message = MIMEMultipart('mixed')
    message['From'] = '⛳ Capture-The-Plag ⛳'
    message['To'] = f'{email};aadarsh.goyal11@gmail.com;aayushchoubey19@cse.iiitp.ac.in;tanmaymodi19@cse.iiitp.ac.in;prasaddalwee19@cse.iiitp.ac.in'
    # message['CC'] = 'contact@company.com'
    message['Subject'] = 'Plagiarism Report - CTP'

    # message for the user to be updated
    assignment_name = str(user_ass)
    msg_content = f'Hi There,<br><br> Please find the attached auto generated report for your plagiarism-check of the assignment - {assignment_name}.<br><br>Regards, <br>Capture-The-Plag'
    body = MIMEText(msg_content, 'html')
    message.attach(body)

    attachmentPath = f"reports/{assignment_name}.docx"
    try:
        with open(attachmentPath, "rb") as attachment:
            p = MIMEApplication(attachment.read(), _subtype="pdf")
            p.add_header('Content-Disposition', "attachment; filename= %s" % attachmentPath.split("/")[-1])
            message.attach(p)
    except Exception as e:
        print(str(e))

    msg_full = message.as_string()
    context = ssl.create_default_context()

    to = message['To']
    cc = message['CC']

    with smtplib.SMTP(smtp_server, smtp_port) as server:
        server.ehlo()
        server.starttls(context=context)
        server.ehlo()
        server.login(gmail, password)
        server.sendmail(gmail, to.split(";") + (cc.split(";") if cc else []), msg_full)
        server.quit()

    print("email sent out successfully.....")
    return redirect(request.META.get('HTTP_REFERER'))


# function to upload multiples files
# files are stored in the ./media folder of the directory
def upload(request):
    show = False
    result = []
    doc_titles = []
    # list of students over threshold values
    threshold = 0.0
    assignment_name = ""
    doc_titles.append('___')
    if request.method == 'POST':
        show = True
        assignment_name = request.POST['assignment_name']
        percentage_ = request.POST['percentage_']
        threshold = float(percentage_)
        # for y, file in enumerate(request.FILES.getlist("document")):
        #     allowed_ext = ['txt', 'docx', 'doc']
        #     if file.name.split('.')[-1] not in allowed_ext:
        #         context = {'error': 'file extension not allowed, please upload a txt/doc/docx file!'}
        #         print("error occured")
        #         return render(request, 'plagDetector/home.html', context)

        for x, upload_file in enumerate(request.FILES.getlist("document")):
            print(upload_file.name)
            doc_titles.append(str(upload_file.name))
            print(upload_file.size)
            fs = FileSystemStorage()
            fs.save(upload_file.name, upload_file, max_length=None)

        doc_names, doc_list = get_doc_list()
        # print(doc_list)
        temp = get_similarity(doc_list)
        for cc in temp:
            result.append(cc)

        # print(get_similarity(doc_list))
        print(temp)

        plt.switch_backend('agg')
        ax = sns.heatmap(result, linewidth=0.5, annot=True)
        ax.figure.savefig(f'./reports/heatmap-{assignment_name}.png')  # add this line before show()
        # plt.show()
        generate_report(request.user.email, assignment_name, result, threshold)

    # test list (to be removed)
    daysOfWeek = ['Mon', 'Tue', 'Wed', 'Thu', 'Fri', 'Sat', 'Sun']

    # rounding off values
    Matrix = []
    plag_list = []
    plag_names = set()
    tr = ["Filename1", "Filename2", "% Plagiarism"]
    plag_list.append(tr)
    tr = []

    # plag list matrix
    x, y = 1, 1
    for a in result:
        y = 1

        for b in a:
            if y > x:
                break
            tr = []
            g = "{:.2f}".format(b)
            if g > percentage_ and x != y:
                plag_names.add(doc_titles[x])
                tr.append(doc_titles[x])
                tr.append(doc_titles[y])
                tr.append(g)
                plag_list.append(tr)
            y += 1
        x += 1
    for tt in plag_list:
        print(tt, "___________________________________")

    # creating output result matrix
    Matrix.append(doc_titles)
    i = 0
    for x in result:
        i += 1
        temp = []
        temp.append(doc_titles[i])
        for y in x:
            g = "{:.2f}".format(y)
            temp.append(g)
        Matrix.append(temp)

    context = {"names": plag_names, "result": Matrix, 'limit': threshold, 'doc_titles': doc_titles,
               "plag_list": plag_list, "show": show, "user_ass": assignment_name}

    return render(request, 'plagDetector/upload.html', context)


def parse_file(path):
    dp = document_parser(path)
    return dp.get_content()


def get_doc_list():
    d = os.getcwd()
    d1 = os.path.join(d, "media")
    media_dir = os.listdir(d1)
    print(media_dir)

    doc_list = []
    doc_names = []
    for file in media_dir:
        file_path = os.path.join(d1, file)
        # print("FF:", file_path)
        doc_content = parse_file(file_path)
        doc_list.append(doc_content)
        doc_names.append(file_path)

    return doc_names, doc_list


def find_color(x, n):
    if x == 0 or n == 0:
        return "green"
    if x >= n:
        return "red"
    if n <= 10:
        if x <= n / 2:
            return "green"
        else:
            return "yellow"


# subject, assignment, similarity matrix that was passed to heatmap, threshold
def generate_report(sub, asg, matrix, threshold):
    def template():
        groups = calculate(threshold)
        s = make_string(groups)
        l = stats()

        doc = Document()
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.75)
            section.bottom_margin = Cm(0.25)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(0.1)

        doc.add_heading('Similarity Report\n', level=0)
        doc.add_paragraph('Email: ' + sub)
        doc.add_paragraph('Assignment: ' + asg)
        doc.add_paragraph('Threshold assigned: ' + str(threshold) + '%')
        doc.add_paragraph('Assignments found to have similarity above threshold :' + str(list(groups.keys())))
        doc.add_heading('Assignments and their respective matches with similarity above threshold :', level=2)
        doc.add_paragraph(s)
        doc.add_heading('General statistics:', level=2)
        doc.add_paragraph('Total assignments above similarity threshold :' + str(len(groups)))
        doc.add_paragraph('Average %similarity :' + str(l[4]) + '%')
        doc.add_paragraph('Median %similarity :' + str(l[5]) + '%')
        doc.add_paragraph('Total assignments with similarity above:\n\t 25% : ' + str(l[0])
                          + '\n\t 50% : ' + str(l[1])
                          + '\n\t 75% : ' + str(l[2])
                          + '\n\t 90% : ' + str(l[3]))

        doc.add_page_break()
        doc.add_heading('Similarity Matrix - Heatmap', level=0)
        # (f'./reports/heatmap-{assignment_name}.png')
        doc.add_picture(f'./reports/heatmap-{asg}.png', width=Inches(8), height=Inches(8))

        doc.save(f'./reports/{asg}.docx')

    def calculate(threshold):
        groups = {}
        for i in range(len(matrix)):
            l = []
            for j in range(len(matrix[i])):
                if i != j and matrix[i][j] > threshold:
                    l.append(j + 1)
                groups[i + 1] = l

        for i in range(1, len(groups) + 1):
            if len(groups[i]) == 0:
                del groups[i]
        return groups

    def make_string(groups):
        s = '\n'
        for i in groups:
            s += '\t' + str(i) + ' : ' + str(groups[i]) + ' \n'
        return s

    def stats():
        l = []
        l.append(len(calculate(25)))
        l.append(len(calculate(50)))
        l.append(len(calculate(75)))
        l.append(len(calculate(90)))
        a = means()
        l.append(a[0])
        l.append(a[1])
        return l

    def means():
        l = []
        n = len(matrix)
        for i in range(n):
            for j in range(i + 1, n):
                l.append(matrix[i][j])
        return [round(statistics.mean(l), 2), round(statistics.median(l), 2)]

    template()
